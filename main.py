"""This Rag system has been built to show the implementation of Task 1 and Task 2 together.

Task 1: Load data from a PDF file, create a FAISS vector store, and run RAG queries using Groq LLM.

Task 2: Extended the system to also load data from TXT and DOCX files, integrate them into the same FAISS vector store, and handle RAG queries that utilize information from all three document types.  

I have build a single RAG system that can handle documents from all three formats (PDF, TXT, DOCX) seamlessly.

To check the individual implementations of Task 1 and Task 2, please refer to the respective files named "main.py" in the same directory.

You can run this code by commenting the files of type you want to check and see if the vector store is created/loaded correctly and the RAG queries work as expected.

To run use command python main.py


"""



import os
import json
import hashlib
from groq import Groq
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq
from docx import Document 
from langchain_core.documents import Document as LCDocument 
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from dotenv import load_dotenv
load_dotenv()


GROQ_API_KEY = os.getenv("GROQ_API_KEY") 

if not GROQ_API_KEY:
    raise ValueError("GROQ_API_KEY environment variable not set. Please set it.")

PDF_PATH = r"yash_resume.pdf"
# To check Task 1 only (PDF), comment out the following lines:
TEXT_PATH = r"suresh_resume.txt" 
DOCX_PATH = r"ramesh_resume.docx"
FAISS_DB_PATH = "vector_db"
FILE_METADATA_PATH = "vector_db/file_metadata.json"
MODEL_NAME = "llama-3.1-8b-instant"
EMBEDDING_MODEL = "all-MiniLM-L6-v2"

def get_file_hash(file_path):
    """Calculate MD5 hash of a file to detect changes"""
    if not os.path.exists(file_path):
        return None
    
    hash_md5 = hashlib.md5()
    try:
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    except Exception as e:
        print(f"Error calculating hash for {file_path}: {e}")
        return None

def load_file_metadata():
    """Load metadata about previously processed files"""
    if os.path.exists(FILE_METADATA_PATH):
        try:
            with open(FILE_METADATA_PATH, 'r') as f:
                return json.load(f)
        except Exception as e:
            print(f"Error loading metadata: {e}")
            return {}
    return {}

def save_file_metadata(metadata):
    """Save metadata about processed files"""
    os.makedirs(os.path.dirname(FILE_METADATA_PATH), exist_ok=True)
    try:
        with open(FILE_METADATA_PATH, 'w') as f:
            json.dump(metadata, f, indent=2)
    except Exception as e:
        print(f"Error saving metadata: {e}")

def detect_file_changes(file_paths):
    """Detect which files are new or modified"""
    metadata = load_file_metadata()
    new_files = []
    modified_files = []
    unchanged_files = []
    
    for file_path in file_paths:
        if not file_path or not os.path.exists(file_path):
            continue
            
        current_hash = get_file_hash(file_path)
        if not current_hash:
            continue
            
        if file_path not in metadata:
            new_files.append(file_path)
            print(f"✨ NEW FILE detected: {file_path}")
        elif metadata[file_path]['hash'] != current_hash:
            modified_files.append(file_path)
            print(f"📝 MODIFIED FILE detected: {file_path}")
        else:
            unchanged_files.append(file_path)
            print(f"✅ UNCHANGED FILE: {file_path}")
    
    return new_files, modified_files, unchanged_files

def extract_text_from_docx(file_path):
    if not os.path.exists(file_path):
        print(f"Warning: Word file not found at {file_path}")
        return ""
    try:
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
        return ""

def load_single_document(file_path):
    """Load a single document and return chunks"""
    documents = []
    
    if not os.path.exists(file_path):
        print(f"Warning: File not found at {file_path}")
        return []
    
    try:
        if file_path.endswith('.pdf'):
            print(f"Loading PDF: {file_path}...")
            loader = PyPDFLoader(file_path)
            documents.extend(loader.load())
            
        elif file_path.endswith('.txt'):
            print(f"Loading TXT: {file_path}...")
            loader = TextLoader(file_path, encoding='utf-8')
            documents.extend(loader.load())
            
        elif file_path.endswith('.docx'):
            print(f"Loading DOCX: {file_path}...")
            docx_text = extract_text_from_docx(file_path)
            if docx_text:
                documents.append(
                    LCDocument(page_content=docx_text, metadata={'source': file_path})
                )
    except Exception as e:
        print(f"Error loading {file_path}: {e}")
        return []
    
    if documents:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=550, 
            chunk_overlap=55,
            length_function=len
        )
        chunks = text_splitter.split_documents(documents)
        print(f"  → Created {len(chunks)} chunks from {file_path}")
        return chunks
    
    return []

def load_and_split_documents(file_paths):
    """Load and split multiple documents"""
    all_chunks = []
    
    for file_path in file_paths:
        if file_path:
            chunks = load_single_document(file_path)
            all_chunks.extend(chunks)
    
    print(f"\n📊 Total chunks created: {len(all_chunks)}")
    return all_chunks

def check_db_exists(db_path):
    return os.path.exists(db_path) and os.path.exists(os.path.join(db_path, "index.faiss"))

def create_vector_store(text_chunks, db_path):
    """Create a new vector store from scratch"""
    print("🔨 Creating NEW FAISS Vector Store...")
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
    vectorstore = FAISS.from_documents(text_chunks, embeddings)
    vectorstore.save_local(db_path)
    print(f"💾 FAISS index saved to {db_path}")
    return vectorstore

def update_vector_store(new_chunks, db_path):
    """Add new chunks to existing vector store"""
    print("🔄 UPDATING existing FAISS Vector Store with new documents...")
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
    
    vectorstore = FAISS.load_local(db_path, embeddings, allow_dangerous_deserialization=True)
    
    vectorstore.add_documents(new_chunks)
    
    vectorstore.save_local(db_path)
    print(f"💾 Updated FAISS index saved to {db_path}")
    return vectorstore

def load_vector_store(db_path):
    """Load existing vector store"""
    print(f"📂 Loading existing FAISS Vector Store from {db_path}...")
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
    return FAISS.load_local(db_path, embeddings, allow_dangerous_deserialization=True)

def update_metadata_for_files(file_paths):
    """Update metadata for processed files"""
    metadata = load_file_metadata()
    
    for file_path in file_paths:
        if file_path and os.path.exists(file_path):
            file_hash = get_file_hash(file_path)
            if file_hash:
                metadata[file_path] = {
                    'hash': file_hash,
                    'last_processed': os.path.getmtime(file_path)
                }
    
    save_file_metadata(metadata)

def format_docs(docs):
    return "\n\n".join([doc.page_content for doc in docs])

def run_rag_query(vectorstore, query, groq_api_key):
    print(f"\n--- Running Query: '{query}' ---")
    
    llm = ChatGroq(
        groq_api_key=groq_api_key, 
        model_name=MODEL_NAME
    )
    
    retriever = vectorstore.as_retriever(search_kwargs={"k": 3}) 
    
    template = """You are a highly knowledgeable expert resume analyst. Your task is to provide detailed and accurate answers.
    
    RULES:
    1. Always use the 'Context' provided below to answer specific questions about the individuals (Yash Raj, Suresh Kumar, Ramesh Kumar).
    2. For general or comparative queries (e.g., "what is a data scientist?"), you may use your extensive external knowledge to provide comprehensive and insightful details.
    3. Be conversational and helpful.

Context (Primary Source):
{context}

Question: {question}

Answer:"""
    
    prompt = ChatPromptTemplate.from_template(template)
    
    rag_chain = (
        {"context": retriever | format_docs, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )

    response = rag_chain.invoke(query)
    source_documents = retriever.invoke(query) 
    
    print("\n✅ GENERATED RESPONSE:")
    print(response)
    print("\n📚 SOURCES USED:")
    sources = set([doc.metadata.get('source') for doc in source_documents])
    for source in sources:
        print(f"- {source}")
    print("-" * 50)


def main():
    print("="*60)
    print("Starting RAG System Initialization...")
    print("="*60)
    
    active_paths = []
    if 'PDF_PATH' in globals() and PDF_PATH:
        active_paths.append(PDF_PATH)
    if 'TEXT_PATH' in globals() and TEXT_PATH:
        active_paths.append(TEXT_PATH)
    if 'DOCX_PATH' in globals() and DOCX_PATH:
        active_paths.append(DOCX_PATH)
    
    print(f"\n📁 Configured files to process: {len(active_paths)}")
    for path in active_paths:
        print(f"  - {path}")
    
    print("\n🔍 Checking for file changes...")
    new_files, modified_files, unchanged_files = detect_file_changes(active_paths)
    
    files_to_process = new_files + modified_files
    
    db_exists = check_db_exists(FAISS_DB_PATH)
    
    if not db_exists:
        print("\n🆕 No existing vector database found. Creating new database...")
        all_chunks = load_and_split_documents(active_paths)
        
        if not all_chunks:
            print("❌ ERROR: No documents could be loaded. Cannot proceed.")
            return
        
        vectorstore = create_vector_store(all_chunks, FAISS_DB_PATH)
        update_metadata_for_files(active_paths)
        
    elif files_to_process:
        print(f"\n🔄 Found {len(files_to_process)} file(s) to process:")
        for f in files_to_process:
            status = "NEW" if f in new_files else "MODIFIED"
            print(f"  - [{status}] {f}")
        
        new_chunks = load_and_split_documents(files_to_process)
        
        if new_chunks:
            vectorstore = update_vector_store(new_chunks, FAISS_DB_PATH)
            update_metadata_for_files(files_to_process)
        else:
            print("⚠️ No new chunks created. Loading existing database.")
            vectorstore = load_vector_store(FAISS_DB_PATH)
    else:
        print("\n✅ All files are up-to-date. Using existing vector database.")
        vectorstore = load_vector_store(FAISS_DB_PATH)
    
    print("\n" + "="*60)
    print("💬 Interactive RAG Chat Mode")
    print("="*60)
    print("The system is ready. Ask questions about your documents.")
    print("Type 'exit' or 'quit' to end.\n")
    
    while True:
        user_query = input("Your Query > ")
        
        if user_query.lower() in ["exit", "quit"]:
            print("\n👋 Exiting RAG chat. Goodbye!")
            break
        
        if not user_query.strip():
            continue
            
        try:
            run_rag_query(vectorstore, user_query, GROQ_API_KEY)
        except Exception as e:
            print(f"❌ An error occurred during query execution: {e}")

if __name__ == "__main__":
    main()