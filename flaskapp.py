"""To run use command python flaskapp.py"""


"""
curl to check the flask APP in postman or terminal


curl -X POST http://127.0.0.1:5000/rag_query \
-H "Content-Type: application/json" \
-d '{"query": "What are the key technical skills of Yash Raj and what is his highest degree?"}'

"""

import os
import json
import hashlib
from groq import Groq
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq
from docx import Document 
from langchain_core.documents import Document as LCDocument 
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from dotenv import load_dotenv
from flask import Flask, request, jsonify

load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY") 

if not GROQ_API_KEY:
    raise ValueError("GROQ_API_KEY environment variable not set. Please set it.")


# To check Task 1 only (PDF), comment out the following lines:

PDF_PATH = r"yash_resume.pdf"
TEXT_PATH = r"suresh_resume.txt" 
DOCX_PATH = r"ramesh_resume.docx"
FAISS_DB_PATH = "flask_vector_db"
FILE_METADATA_PATH = "flask_vector_db/file_metadata.json"
MODEL_NAME = "llama-3.1-8b-instant"
EMBEDDING_MODEL = "all-MiniLM-L6-v2"

app = Flask(__name__)

def get_file_hash(file_path):
    if not os.path.exists(file_path):
        return None
    
    hash_md5 = hashlib.md5()
    try:
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    except Exception as e:
        print(f"Error calculating hash for {file_path}: {e}")
        return None

def load_file_metadata():
    if os.path.exists(FILE_METADATA_PATH):
        try:
            with open(FILE_METADATA_PATH, 'r') as f:
                return json.load(f)
        except Exception as e:
            print(f"Error loading metadata: {e}")
            return {}
    return {}

def save_file_metadata(metadata):
    os.makedirs(os.path.dirname(FILE_METADATA_PATH), exist_ok=True)
    try:
        with open(FILE_METADATA_PATH, 'w') as f:
            json.dump(metadata, f, indent=2)
    except Exception as e:
        print(f"Error saving metadata: {e}")

def detect_file_changes(file_paths):
    metadata = load_file_metadata()
    new_files = []
    modified_files = []
    unchanged_files = []
    
    for file_path in file_paths:
        if not file_path or not os.path.exists(file_path):
            continue
            
        current_hash = get_file_hash(file_path)
        if not current_hash:
            continue
            
        if file_path not in metadata:
            new_files.append(file_path)
        elif metadata[file_path]['hash'] != current_hash:
            modified_files.append(file_path)
        else:
            unchanged_files.append(file_path)
    
    return new_files, modified_files, unchanged_files

def extract_text_from_docx(file_path):
    if not os.path.exists(file_path):
        return ""
    try:
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
        return ""

def load_single_document(file_path):
    documents = []
    
    if not os.path.exists(file_path):
        return []
    
    try:
        if file_path.endswith('.pdf'):
            loader = PyPDFLoader(file_path)
            documents.extend(loader.load())
            
        elif file_path.endswith('.txt'):
            loader = TextLoader(file_path, encoding='utf-8')
            documents.extend(loader.load())
            
        elif file_path.endswith('.docx'):
            docx_text = extract_text_from_docx(file_path)
            if docx_text:
                documents.append(
                    LCDocument(page_content=docx_text, metadata={'source': file_path})
                )
    except Exception as e:
        print(f"Error loading {file_path}: {e}")
        return []
    
    if documents:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=550, 
            chunk_overlap=55,
            length_function=len
        )
        chunks = text_splitter.split_documents(documents)
        return chunks
    
    return []

def load_and_split_documents(file_paths):
    all_chunks = []
    
    for file_path in file_paths:
        if file_path:
            chunks = load_single_document(file_path)
            all_chunks.extend(chunks)
    
    return all_chunks

def check_db_exists(db_path):
    return os.path.exists(db_path) and os.path.exists(os.path.join(db_path, "index.faiss"))

def create_vector_store(text_chunks, db_path):
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
    vectorstore = FAISS.from_documents(text_chunks, embeddings)
    vectorstore.save_local(db_path)
    return vectorstore

def update_vector_store(new_chunks, db_path):
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
    
    vectorstore = FAISS.load_local(db_path, embeddings, allow_dangerous_deserialization=True)
    
    vectorstore.add_documents(new_chunks)
    
    vectorstore.save_local(db_path)
    return vectorstore

def load_vector_store(db_path):
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
    return FAISS.load_local(db_path, embeddings, allow_dangerous_deserialization=True)

def update_metadata_for_files(file_paths):
    metadata = load_file_metadata()
    
    for file_path in file_paths:
        if file_path and os.path.exists(file_path):
            file_hash = get_file_hash(file_path)
            if file_hash:
                metadata[file_path] = {
                    'hash': file_hash,
                    'last_processed': os.path.getmtime(file_path)
                }
    
    save_file_metadata(metadata)

def format_docs(docs):
    return "\n\n".join([doc.page_content for doc in docs])

def initialize_vector_store():
    active_paths = []
    if 'PDF_PATH' in globals() and PDF_PATH:
        active_paths.append(PDF_PATH)
    if 'TEXT_PATH' in globals() and TEXT_PATH:
        active_paths.append(TEXT_PATH)
    if 'DOCX_PATH' in globals() and DOCX_PATH:
        active_paths.append(DOCX_PATH)
    
    new_files, modified_files, unchanged_files = detect_file_changes(active_paths)
    files_to_process = new_files + modified_files
    db_exists = check_db_exists(FAISS_DB_PATH)

    if not db_exists:
        all_chunks = load_and_split_documents(active_paths)
        if not all_chunks:
            return None
        vectorstore = create_vector_store(all_chunks, FAISS_DB_PATH)
        update_metadata_for_files(active_paths)
    elif files_to_process:
        new_chunks = load_and_split_documents(files_to_process)
        if new_chunks:
            vectorstore = update_vector_store(new_chunks, FAISS_DB_PATH)
            update_metadata_for_files(files_to_process)
        else:
            vectorstore = load_vector_store(FAISS_DB_PATH)
    else:
        vectorstore = load_vector_store(FAISS_DB_PATH)
    
    return vectorstore

vectorstore = initialize_vector_store()

@app.route("/rag_query", methods=["POST"])
def rag_query():
    data = request.get_json()
    query = data.get("query")

    if not query:
        return jsonify({"error": "Missing 'query' parameter"}), 400

    if not vectorstore:
        return jsonify({"error": "Vector store initialization failed."}), 500

    llm = ChatGroq(
        groq_api_key=GROQ_API_KEY, 
        model_name=MODEL_NAME
    )
    
    retriever = vectorstore.as_retriever(search_kwargs={"k": 3}) 
    
    template = """You are a highly knowledgeable expert resume analyst. Your task is to provide detailed and accurate answers.
    
    RULES:
    1. Always use the 'Context' provided below to answer specific questions about the individuals (Yash Raj, Suresh Kumar, Ramesh Kumar).
    2. For general or comparative queries (e.g., "what is a data scientist?"), you may use your extensive external knowledge to provide comprehensive and insightful details.
    3. Be conversational and helpful.

Context (Primary Source):
{context}

Question: {question}

Answer:"""
    
    prompt = ChatPromptTemplate.from_template(template)
    
    rag_chain = (
        {"context": retriever | format_docs, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )

    try:
        response = rag_chain.invoke(query)
        source_documents = retriever.invoke(query) 
        
        sources = list(set([doc.metadata.get('source') for doc in source_documents]))
        
        return jsonify({
            "query": query,
            "response": response,
            "sources": sources
        })
    except Exception as e:
        print(f"Error during query execution: {e}")
        return jsonify({"error": f"An internal error occurred: {e}"}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)